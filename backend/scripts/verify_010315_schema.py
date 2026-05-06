"""verify_010315_schema.py — 驗證 010315 schema 的 marker / cell 在實際 docx 中找得到。

執行：
    cd backend && python scripts/verify_010315_schema.py

檢查項目：
- 每個 para field：para_idx 存在，且 marker 字串確實在該 paragraph text 內
- 每個 cell field：(table_idx, row, col) 存在
- 每個 cell_marker field：cell 存在，且 marker 字串在某個 paragraph 內
- 統計命中／未命中數，未命中印詳細位置與該 cell/para 的實際內容（前 60 字）
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

SCHEMA_PATH = Path(__file__).parent.parent / "app" / "rag" / "form_schemas" / "010315.json"
DOCX_PATH = (Path(__file__).parent.parent.parent / "data_markdown"
             / "form_data" / "010315工地文件管制與保存表.docx")


def show(s: str, n: int = 60) -> str:
    """顯示用：把 \\t / \\n 視覺化，方便看出 marker 對齊。"""
    s = s.replace("\t", "↹").replace("\n", "⏎")
    return s if len(s) <= n else s[:n] + "…"


def main() -> None:
    schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
    doc = Document(str(DOCX_PATH))

    ok = 0
    fails: list[str] = []

    for f in schema["fields"]:
        loc = f["loc"]
        kind = loc["kind"]
        key = f["key"]

        if kind == "para":
            idx = loc["para_idx"]
            if idx >= len(doc.paragraphs):
                fails.append(f"[{key}] para_idx {idx} 超出範圍（共 {len(doc.paragraphs)} 段）")
                continue
            text = doc.paragraphs[idx].text
            marker = loc["marker"]
            if marker in text:
                ok += 1
            else:
                fails.append(
                    f"[{key}] para[{idx}] 找不到 marker={show(marker, 30)!r}\n"
                    f"     段落實際: {show(text)!r}"
                )

        elif kind in ("cell", "cell_marker"):
            ti, ri, ci = loc["table_idx"], loc["row"], loc["col"]
            if ti >= len(doc.tables):
                fails.append(f"[{key}] table_idx {ti} 超出範圍")
                continue
            table = doc.tables[ti]
            if ri >= len(table.rows) or ci >= len(table.columns):
                fails.append(f"[{key}] tbl{ti} r{ri}c{ci} 超出範圍 ({len(table.rows)} 列 × {len(table.columns)} 欄)")
                continue
            cell = table.rows[ri].cells[ci]
            if kind == "cell":
                ok += 1  # 純 cell 寫入不需 marker，loc 存在即合法
            else:  # cell_marker
                marker = loc["marker"]
                hit = any(marker in p.text for p in cell.paragraphs)
                if hit:
                    ok += 1
                else:
                    fails.append(
                        f"[{key}] tbl{ti} r{ri}c{ci} 找不到 marker={show(marker, 30)!r}\n"
                        f"     cell 實際: {show(cell.text)!r}"
                    )
        else:
            fails.append(f"[{key}] 未知 kind={kind!r}")

    total = len(schema["fields"])
    print(f"總欄位 {total}    通過 {ok}    失敗 {len(fails)}")
    if fails:
        print("\n=== 失敗項目 ===")
        for line in fails:
            print(line)
            print()


if __name__ == "__main__":
    main()
