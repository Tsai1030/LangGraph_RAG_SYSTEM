"""inspect_form.py — 偵錯用：dump 一份 .docx 的 paragraph 與 table 真實結構。

用法：
    cd backend && python scripts/inspect_form.py [檔名.docx]

不傳檔名時預設 dump 010315。產出在 scripts/output/inspect_<stem>.txt。

dump 內容：
- 所有非空段落（index、style、文字）— 用來找章節 heading（附件 1/2/3 等）
- 每張 table 的尺寸與每個 cell 的內容，並標出水平/垂直合併、空 cell。

合併偵測：
- python-docx 的 row.cells 會把合併儲存格展開為多個 cell，重複的 cell 共用 _tc element。
- 同 row 內出現第 2+ 次的 _tc → 水平合併重複位置（標 ←H）
- vMerge=continue 屬性 → 垂直合併延續列（標 ↑V）
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

FORMS_DIR = Path(__file__).parent.parent.parent / "data_markdown" / "form_data"
OUT_DIR = Path(__file__).parent / "output"

DEFAULT_FILENAME = "010315工地文件管制與保存表.docx"


def _clean(s: str) -> str:
    return s.replace("\t", "↹").replace("\n", "⏎").strip()


def _is_vmerge_continue(tc) -> bool:
    """偵測 cell 的 <w:vMerge> 是否為 continue（無 val 或 val=continue）。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    vmerge = tcPr.find(qn("w:vMerge"))
    if vmerge is None:
        return False
    val = vmerge.get(qn("w:val"))
    return val is None or val == "continue"


def dump(filename: str) -> Path:
    path = FORMS_DIR / filename
    if not path.exists():
        candidates = sorted(p.name for p in FORMS_DIR.glob("*.docx") if not p.name.startswith("~"))
        raise FileNotFoundError(
            f"找不到 {path}\n可用檔案：\n  - " + "\n  - ".join(candidates)
        )

    doc = Document(path)
    OUT_DIR.mkdir(exist_ok=True)
    out = OUT_DIR / f"inspect_{path.stem}.txt"

    lines: list[str] = []
    lines.append(f"檔案：{path.name}")
    lines.append(f"段落數：{len(doc.paragraphs)}")
    lines.append(f"表格數：{len(doc.tables)}")
    lines.append("")

    # ── 段落（heading + 非空 normal）─────────────────────────────────
    lines.append("=" * 80)
    lines.append("段落（paragraphs）— 只列非空 + heading 風格")
    lines.append("=" * 80)
    for i, p in enumerate(doc.paragraphs):
        text = _clean(p.text)
        style = p.style.name if p.style else ""
        is_heading = "Heading" in style or "Title" in style
        if not text and not is_heading:
            continue
        marker = "★" if is_heading else " "
        lines.append(f"  {marker}[{i:3d}] style={style:<20} text={text!r}")
    lines.append("")

    # ── 表格（每個 cell + 合併）─────────────────────────────────────
    for ti, table in enumerate(doc.tables):
        lines.append("=" * 80)
        lines.append(f"表 {ti}  ({len(table.rows)} 列 × {len(table.columns)} 欄)")
        lines.append("=" * 80)

        for ri, row in enumerate(table.rows):
            cells = row.cells
            seen_tc_ids: list[int] = []
            for ci, cell in enumerate(cells):
                tc = cell._tc
                tc_id = id(tc)

                tags: list[str] = []
                if tc_id in seen_tc_ids:
                    tags.append("←H")  # 同 row 水平合併重複位
                else:
                    seen_tc_ids.append(tc_id)
                if _is_vmerge_continue(tc):
                    tags.append("↑V")  # 垂直合併延續

                text = _clean(cell.text)
                if not text and not tags:
                    tags.append("（空）")

                tag_str = " ".join(tags)
                lines.append(f"  r{ri:02d}c{ci:02d}: {text!r:<55}  {tag_str}")
            lines.append("")

    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def main() -> None:
    filename = sys.argv[1] if len(sys.argv) >= 2 else DEFAULT_FILENAME
    out = dump(filename)
    size = out.stat().st_size
    print(f"OK → {out}  ({size:,} bytes)")


if __name__ == "__main__":
    main()
