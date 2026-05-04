"""
build_form_schemas.py — 離線解析靜態 .docx 表單，產出欄位 schema JSON。

執行：
    cd backend && python scripts/build_form_schemas.py

產出：
    backend/app/rag/form_schemas/<form_id>.json   （三份）
    backend/scripts/output/form_schemas_summary.txt  （人工校對用，列出所有 fields）

schema 結構：
{
  "form_id": "010101",
  "title": "動員開工作業檢核表",
  "file_name": "010101動員開工作業檢核表.docx",
  "fields": [
    {
      "key": "工程名稱",
      "label": "工程名稱",
      "type": "text",
      "required": true,
      "loc": {"kind": "para", "para_idx": 1, "marker": "工程名稱:", "marker_end": "\t"}
    },
    {
      "key": "tbl0_r2_status",
      "label": "編號 2.1「組織提報」— 完成狀態",
      "type": "checkbox_vx",
      "required": false,
      "loc": {"kind": "cell", "table_idx": 0, "row": 2, "col": 7}
    },
    ...
  ]
}

loc 兩種型態：
  - {"kind": "para", "para_idx": int, "marker": "標籤:", "marker_end": "\t" | null}
      → fill 階段：在段落內找到 marker，將值插在 marker 後（marker_end 之前）
  - {"kind": "cell", "table_idx": int, "row": int, "col": int}
      → fill 階段：直接寫入該 cell
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table

FORMS_DIR = Path(__file__).parent.parent.parent / "data_markdown" / "form_data"
REGISTRY_PATH = Path(__file__).parent.parent / "app" / "rag" / "form_registry.json"
OUT_DIR = Path(__file__).parent.parent / "app" / "rag" / "form_schemas"
SUMMARY_PATH = Path(__file__).parent / "output" / "form_schemas_summary.txt"

# 9 欄檢核表的標準欄位（010101 / 010102）
_CHECKLIST_COL_NAMES = ["項次", "項目", "編號", "作業內容", "辦理期限", "主辦單位", "協辦單位", "完成狀態", "備註"]
_CHECKLIST_FILLABLE_COLS = [7, 8]  # 完成狀態 + 備註
_STATUS_COL = 7
_REMARK_COL = 8

# 段落抬頭標籤偵測：以 \t 切片後，每片需形如 'XX:' 或 'XX：' 且冒號後為空
_PARA_CHUNK_RE = re.compile(r"^([^\s:：][^:：]{0,15})[:：]\s*$")

# 視為日期類型的關鍵字
_DATE_HINTS = ("日期", "年月日", "時間")


def _clean(text: str) -> str:
    return text.replace("　", " ").replace("\t", " ").replace("\n", " ").strip()


def _short(text: str, n: int = 30) -> str:
    text = _clean(text).replace("\n", " ").replace("\t", " ")
    return text if len(text) <= n else text[:n] + "…"


def _field_type(label: str) -> str:
    return "date" if any(h in label for h in _DATE_HINTS) else "text"


# ── 段落抬頭欄位 ─────────────────────────────────────────────────────

def parse_header_fields(doc: DocxDocument) -> list[dict[str, Any]]:
    """
    掃描段落，先以 \\t 切片，再判斷每片是否形如 'XX:'（冒號後為空）。
    切片做法可避免像 '單\\t位：' 這種被 tab 切斷的 label 誤抓到單字 '位'。
    例：
      '工程名稱:\\t工令:' → ['工程名稱:', '工令:'] → 兩欄位
      '作業流程第 010101 節\\t製表日期:' → 前段不匹配，'製表日期:' 匹配
    過濾：label 長度 <= 1 的單字一律跳過（多為被 tab 切斷的殘片）。
    """
    fields: list[dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text or (":" not in text and "：" not in text):
            continue
        chunks = text.split("\t")
        for ci, chunk in enumerate(chunks):
            chunk_clean = _clean(chunk)
            m = _PARA_CHUNK_RE.match(chunk_clean)
            if not m:
                continue
            label = _clean(m.group(1))
            if len(label) <= 1:
                continue  # 'XX:' 但 XX 是單字 → 多半是被 tab 切壞
            marker = chunk_clean  # 含冒號
            marker_end = "\t" if ci < len(chunks) - 1 else None
            fields.append({
                "key": label,
                "label": label,
                "type": _field_type(label),
                "required": True,
                "loc": {
                    "kind": "para",
                    "para_idx": idx,
                    "marker": marker,
                    "marker_end": marker_end,
                },
            })
    return fields


# ── 檢核表（010101 / 010102 樣式）────────────────────────────────────

def _is_checklist_table(table: Table) -> bool:
    if len(table.columns) != 9 or len(table.rows) < 2:
        return False
    header_text = " ".join(_clean(c.text) for c in table.rows[0].cells)
    return ("項次" in header_text and "編號" in header_text
            and "作業內容" in header_text and "完成" in header_text)


def parse_checklist_table(table: Table, table_idx: int) -> list[dict[str, Any]]:
    """
    9 欄檢核表：每個有「編號」的資料列產生兩個 field（完成狀態 + 備註）。
    使用合併後的 cell.text 取編號與作業內容供 label 描述。
    """
    fields: list[dict[str, Any]] = []
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue
        cells = [_clean(c.text) for c in row.cells]
        code = cells[2] if len(cells) > 2 else ""
        action = cells[3] if len(cells) > 3 else ""
        if not code:
            continue  # 該列可能是分組標題

        prefix = f"編號 {code}「{_short(action, 30)}」"

        # 完成狀態（V/X）
        if not cells[_STATUS_COL]:
            fields.append({
                "key": f"tbl{table_idx}_r{r_idx}_status",
                "label": f"{prefix} — 完成狀態",
                "type": "checkbox_vx",
                "required": False,
                "loc": {"kind": "cell", "table_idx": table_idx, "row": r_idx, "col": _STATUS_COL},
            })
        # 備註
        if not cells[_REMARK_COL]:
            fields.append({
                "key": f"tbl{table_idx}_r{r_idx}_remark",
                "label": f"{prefix} — 備註",
                "type": "text",
                "required": False,
                "loc": {"kind": "cell", "table_idx": table_idx, "row": r_idx, "col": _REMARK_COL},
            })
    return fields


# ── 通用表格（010315 樣式）──────────────────────────────────────────

def parse_generic_table(table: Table, table_idx: int) -> list[dict[str, Any]]:
    """
    通用策略：
      1. 找「最像 header 的列」（前 3 列中非空儲存格最多者）作為 column header
      2. 該列以下：用 column header 命名空格欄位
      3. 該列以上：以該列首個非空儲存格作 row_label，命名同列其他空格

    合併儲存格：python-docx row.cells 會展開成所有欄位，相鄰 cell 共用 _tc 表示水平合併，跳過重複。
    """
    fields: list[dict[str, Any]] = []
    rows = table.rows
    if not rows:
        return fields

    n_cols = len(rows[0].cells)
    header_idx = 0
    best_filled = -1
    for r_idx in range(min(3, len(rows))):
        filled = sum(1 for c in rows[r_idx].cells if _clean(c.text))
        if filled > best_filled:
            best_filled = filled
            header_idx = r_idx
    headers = [_clean(c.text) or f"col{i}" for i, c in enumerate(rows[header_idx].cells)]

    for r_idx, row in enumerate(rows):
        if r_idx == header_idx:
            continue

        # 在 header 之前的列：取首個非空 cell 當該列 label
        row_label: str | None = None
        if r_idx < header_idx:
            for cell in row.cells:
                txt = _clean(cell.text)
                if txt:
                    row_label = txt
                    break

        for c_idx, cell in enumerate(row.cells):
            text = _clean(cell.text)
            if text:
                continue
            # 跳過水平合併儲存格的重複位置
            try:
                if c_idx > 0 and row.cells[c_idx - 1]._tc is cell._tc:
                    continue
            except Exception:
                pass

            if r_idx < header_idx and row_label:
                label = f"表 {table_idx + 1}・{row_label}"
                type_hint = row_label
            else:
                col_label = headers[c_idx] if c_idx < len(headers) else f"col{c_idx}"
                label = f"表 {table_idx + 1}・第 {r_idx} 列・{col_label}"
                type_hint = col_label

            fields.append({
                "key": f"tbl{table_idx}_r{r_idx}_c{c_idx}",
                "label": label,
                "type": _field_type(type_hint),
                "required": False,
                "loc": {"kind": "cell", "table_idx": table_idx, "row": r_idx, "col": c_idx},
            })
    return fields


# ── 主流程 ──────────────────────────────────────────────────────────

def build_schema(form: dict[str, Any], path: Path) -> dict[str, Any]:
    doc = Document(path)
    fields: list[dict[str, Any]] = []
    fields.extend(parse_header_fields(doc))
    for ti, table in enumerate(doc.tables):
        if _is_checklist_table(table):
            fields.extend(parse_checklist_table(table, ti))
        else:
            fields.extend(parse_generic_table(table, ti))

    return {
        "form_id": form["form_id"],
        "title": form["display_name"],
        "file_name": form["file_name"],
        "fields": fields,
    }


def write_summary(schemas: list[dict[str, Any]]) -> None:
    SUMMARY_PATH.parent.mkdir(exist_ok=True)
    buf = io.StringIO()
    for s in schemas:
        buf.write(f"\n{'=' * 70}\n")
        buf.write(f"{s['form_id']}  {s['title']}\n")
        buf.write(f"  檔案：{s['file_name']}\n")
        buf.write(f"  共 {len(s['fields'])} 個欄位\n")
        buf.write(f"{'=' * 70}\n")
        for f in s["fields"]:
            loc = f["loc"]
            if loc["kind"] == "para":
                loc_str = f"para[{loc['para_idx']}] marker={loc['marker']!r}"
            else:
                loc_str = f"tbl[{loc['table_idx']}] r{loc['row']}c{loc['col']}"
            req = "*" if f["required"] else " "
            buf.write(f"  {req} {f['key']:<30} ({f['type']:<12}) {loc_str}\n")
            buf.write(f"      label: {f['label']}\n")
    SUMMARY_PATH.write_text(buf.getvalue(), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    registry = json.loads(REGISTRY_PATH.read_text(encoding="utf-8"))

    schemas: list[dict[str, Any]] = []
    for form in registry:
        path = FORMS_DIR / form["file_name"]
        if not path.exists():
            print(f"!! 找不到檔案：{path}")
            continue
        schema = build_schema(form, path)
        out_path = OUT_DIR / f"{form['form_id']}.json"
        out_path.write_text(
            json.dumps(schema, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        schemas.append(schema)
        print(f"OK  {form['form_id']}: {len(schema['fields'])} fields → {out_path.name}")

    write_summary(schemas)
    print(f"\nsummary → {SUMMARY_PATH}")


if __name__ == "__main__":
    main()
