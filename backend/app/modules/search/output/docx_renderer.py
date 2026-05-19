"""python-docx renderer.

Key design choices:
  - Placeholders are `{{slot_key}}` literal strings inside a single run
    each (build_template.py guarantees this convention).
  - When a placeholder has confidence != "high", we paint that run red
    so reviewers immediately spot fields that need verification.
  - Multi-run placeholders ARE supported defensively: we first join all
    runs in a paragraph, do replacement, then rebuild runs. This costs
    formatting on those paragraphs — but our generator avoids that case.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Mapping

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .base import DocumentRenderer

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
_LOW_CONF_COLOR = RGBColor(0xC0, 0x39, 0x2B)   # red
_MISSING_COLOR = RGBColor(0x99, 0x99, 0x99)    # grey


class DocxRenderer(DocumentRenderer):
    def __init__(self, template_path: Path) -> None:
        self._template_path = template_path

    def render(
        self,
        slot_values: Mapping[str, str],
        output_path: Path,
        confidence: Mapping[str, str] | None = None,
    ) -> None:
        if not self._template_path.exists():
            raise FileNotFoundError(
                f"Word template missing: {self._template_path}. "
                "Run `uv run python scripts/build_template.py` to generate it."
            )
        doc = Document(str(self._template_path))
        conf = confidence or {}

        for para in doc.paragraphs:
            self._replace_in_paragraph(para, slot_values, conf)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, slot_values, conf)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    # ── internals ───────────────────────────────────

    def _replace_in_paragraph(
        self,
        para: Paragraph,
        slot_values: Mapping[str, str],
        conf: Mapping[str, str],
    ) -> None:
        if "{{" not in para.text:
            return

        # Fast path: each placeholder already lives in its own run, so we
        # can replace at run granularity and preserve every other run's
        # formatting verbatim.
        run_keys: list[tuple[Run, str] | None] = []
        any_run_match = False
        for run in para.runs:
            m = _PLACEHOLDER_RE.fullmatch(run.text.strip())
            if m:
                run_keys.append((run, m.group(1)))
                any_run_match = True
            else:
                run_keys.append(None)

        if any_run_match:
            for entry in run_keys:
                if entry is None:
                    continue
                run, key = entry
                value = slot_values.get(key, run.text)
                run.text = value
                self._maybe_color(run, key, value, conf)
            # Also handle any stragglers that span run boundaries
            if "{{" in para.text:
                self._slow_path_replace(para, slot_values, conf)
            return

        # Slow path: placeholders span run boundaries (e.g. user edited the
        # template by hand and Word split the run). Collapse + rebuild.
        self._slow_path_replace(para, slot_values, conf)

    def _slow_path_replace(
        self,
        para: Paragraph,
        slot_values: Mapping[str, str],
        conf: Mapping[str, str],
    ) -> None:
        text = para.text
        new_text = _PLACEHOLDER_RE.sub(
            lambda m: slot_values.get(m.group(1), m.group(0)), text
        )
        if new_text == text:
            return
        # Clear all runs except first, then set first to combined text
        for r in para.runs[1:]:
            r.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)

    def _maybe_color(
        self, run: Run, key: str, value: str, conf: Mapping[str, str]
    ) -> None:
        c = conf.get(key, "high")
        if value in ("—", "-"):
            run.font.color.rgb = _MISSING_COLOR
            return
        if c == "low":
            run.font.color.rgb = _LOW_CONF_COLOR
            run.font.bold = True
        elif c == "medium":
            run.font.color.rgb = _LOW_CONF_COLOR
