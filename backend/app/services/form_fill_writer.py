"""
form_fill_writer.py — 將收集到的欄位值寫入 .docx 模板副本。

設計重點：
- 段落欄位：用 regex 取代 marker 後既有值（idempotent，重填會覆蓋而非追加）
- 表格欄位：直接覆寫 cell.text；checkbox 類型映射成 V / X 字元
- 不修改原模板，副本存到 GENERATED_DIR，filename 含 conversation_id 與 timestamp
"""

from __future__ import annotations

import json
import logging
import re
import time
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document

from app.rag.form_lookup import get_form_path

logger = logging.getLogger(__name__)

SCHEMAS_DIR = Path(__file__).parent.parent / "rag" / "form_schemas"
GENERATED_DIR = Path(__file__).parent.parent.parent / "data" / "generated_forms"

# checkbox 自然語言 → V / X
_AFFIRMATIVE = {
    "v", "✓", "✔", "○", "o", "yes", "y", "ok", "true", "1",
    "完成", "已完成", "做完", "已辦", "已做", "對", "是",
}
_NEGATIVE = {
    "x", "✗", "✘", "×", "no", "n", "false", "0",
    "未完成", "未辦", "沒做", "未做", "否", "錯",
}


# ──────────────────────────────────────────────────────────────────
# Schema 載入
# ──────────────────────────────────────────────────────────────────

def load_schema(form_id: str) -> dict[str, Any] | None:
    path = SCHEMAS_DIR / f"{form_id}.json"
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


# ──────────────────────────────────────────────────────────────────
# 值正規化
# ──────────────────────────────────────────────────────────────────

def _normalize_value(value: Any, field_type: str) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    if field_type == "checkbox_vx":
        low = text.lower()
        if low in _AFFIRMATIVE:
            return "V"
        if low in _NEGATIVE:
            return "X"
        # 含關鍵字判斷（例如「已完成」「未完成」）
        if any(k in text for k in ("完成", "已", "做")) and "未" not in text:
            return "V"
        if any(k in text for k in ("未", "沒", "尚未")):
            return "X"
        return text  # 不認得，原樣保留
    return text


# ──────────────────────────────────────────────────────────────────
# 段落填入：以 regex 取代 marker 後的既有值，達到 idempotent
# ──────────────────────────────────────────────────────────────────

def _replace_marker_in_text(text: str, marker: str, marker_end: str | None, value: str) -> str:
    """在 text 內把「marker + 舊值（+ marker_end）」取代為「marker 新值（+ marker_end）」。

    - marker_end=None：吃到 \\t 或行尾為止（最常見場景，如「工程名稱：xxx\\t」）
    - marker_end 指定：marker 與 marker_end 中間的內容整段取代為 value，保留 marker_end
      （適用於跨 \\t 區段的場景，如段落「年\\t月\\t日」用 marker='年' marker_end='日'）
    """
    if marker_end:
        pattern = re.compile(re.escape(marker) + r"[^\n]*?" + re.escape(marker_end))
        return pattern.sub(f"{marker} {value}{marker_end}", text, count=1)
    pattern = re.compile(re.escape(marker) + r"[^\t\n]*")
    return pattern.sub(f"{marker} {value}", text, count=1)


def _fill_paragraphs(doc, schema: dict, collected: dict[str, str]) -> int:
    """
    將段落類型的欄位值寫入。同一段落內的多個 marker 一次處理完再 rewrite。
    """
    para_groups: dict[int, list[dict]] = defaultdict(list)
    for f in schema["fields"]:
        if f["loc"]["kind"] != "para":
            continue
        if f["key"] not in collected or collected[f["key"]] in (None, ""):
            continue
        para_groups[f["loc"]["para_idx"]].append(f)

    written = 0
    for para_idx, fields in para_groups.items():
        if para_idx >= len(doc.paragraphs):
            logger.warning("[form_fill] paragraph idx %d out of range", para_idx)
            continue
        para = doc.paragraphs[para_idx]
        text = para.text
        original = text
        for f in fields:
            marker = f["loc"]["marker"]
            marker_end = f["loc"].get("marker_end")
            value = _normalize_value(collected[f["key"]], f["type"])
            new_text = _replace_marker_in_text(text, marker, marker_end, value)
            if new_text != text:
                text = new_text
                written += 1
        if text == original:
            continue
        # 重寫段落內容（保留第一個 run 的格式，其他清空）
        if not para.runs:
            para.add_run(text)
        else:
            para.runs[0].text = text
            for run in para.runs[1:]:
                run.text = ""
    return written


# ──────────────────────────────────────────────────────────────────
# 表格填入
# ──────────────────────────────────────────────────────────────────

def _write_cell_marker(cell, marker: str, marker_end: str | None, value: str) -> bool:
    """在 cell 內找到含 marker 的 paragraph，取代 marker 後的內容為 value。

    一個 cell 可能含多個 paragraph；只取代第一個含 marker 的。
    保留第一個 run 的格式。
    """
    for para in cell.paragraphs:
        text = para.text
        if marker not in text:
            continue
        new_text = _replace_marker_in_text(text, marker, marker_end, value)
        if new_text == text:
            return False
        if not para.runs:
            para.add_run(new_text)
        else:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        return True
    return False


def _fill_cells(doc, schema: dict, collected: dict[str, str]) -> int:
    written = 0
    for f in schema["fields"]:
        kind = f["loc"]["kind"]
        if kind not in ("cell", "cell_marker"):
            continue
        if f["key"] not in collected or collected[f["key"]] in (None, ""):
            continue
        loc = f["loc"]
        try:
            table = doc.tables[loc["table_idx"]]
            cell = table.rows[loc["row"]].cells[loc["col"]]
        except (IndexError, KeyError):
            logger.warning("[form_fill] cell loc out of range: %s", loc)
            continue
        value = _normalize_value(collected[f["key"]], f["type"])
        if kind == "cell":
            cell.text = value
            written += 1
        else:  # cell_marker
            if _write_cell_marker(cell, loc["marker"], loc.get("marker_end"), value):
                written += 1
            else:
                logger.warning("[form_fill] cell_marker not found in cell: %s", loc)
    return written


# ──────────────────────────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────────────────────────

def write_filled_docx(
    form_id: str,
    collected: dict[str, str],
    conversation_id: str,
) -> tuple[Path, int] | None:
    """
    開啟模板副本、寫入欄位、存到 GENERATED_DIR。
    回傳 (output_path, written_count) 或 None（找不到模板/schema）。
    """
    template_path = get_form_path(form_id)
    if template_path is None or not template_path.exists():
        logger.warning("[form_fill] template not found: %s", form_id)
        return None
    schema = load_schema(form_id)
    if schema is None:
        logger.warning("[form_fill] schema not found: %s", form_id)
        return None

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template_path))
    written = _fill_paragraphs(doc, schema, collected) + _fill_cells(doc, schema, collected)

    out_name = f"{conversation_id}_{form_id}_{int(time.time())}.docx"
    out_path = GENERATED_DIR / out_name
    doc.save(str(out_path))
    logger.info("[form_fill] wrote %d fields to %s", written, out_path.name)
    return out_path, written


def get_filled_path(token: str) -> Path | None:
    """根據 token（filename）回傳 generated 檔案路徑。"""
    # 防護：不允許路徑分隔
    if "/" in token or "\\" in token or ".." in token:
        return None
    path = GENERATED_DIR / token
    return path if path.is_file() else None


def delete_generated_for_conversation(conversation_id: str) -> int:
    """
    刪除 GENERATED_DIR 下所有以 <conversation_id>_ 開頭的產出檔
    （含靜態填寫 .docx 與動態匯出 .xlsx / .csv）。
    回傳成功刪除的檔案數；單檔失敗會記 log 但不中斷。

    路徑穿越防護：conversation_id 不接受任何 path separator 或父層引用，
    避免「..」之類的字串造成 glob 跨目錄。
    """
    if not conversation_id or any(c in conversation_id for c in ("/", "\\", "..")):
        logger.warning("[form_fill] reject conversation_id with path-like chars: %r", conversation_id)
        return 0

    if not GENERATED_DIR.exists():
        return 0

    deleted = 0
    gen_resolved = GENERATED_DIR.resolve()
    for path in GENERATED_DIR.glob(f"{conversation_id}_*"):
        if not path.is_file():
            continue
        # 二次驗證：解析後仍位於 GENERATED_DIR 內
        try:
            resolved = path.resolve()
            if gen_resolved not in resolved.parents:
                logger.warning("[form_fill] skip path outside GENERATED_DIR: %s", resolved)
                continue
            path.unlink()
            deleted += 1
        except OSError as exc:
            logger.warning("[form_fill] failed to delete %s: %s", path.name, exc)
    if deleted:
        logger.info("[form_fill] deleted %d generated files for conversation %s", deleted, conversation_id)
    return deleted
